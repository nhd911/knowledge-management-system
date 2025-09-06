from openai import OpenAI
from docx import Document
from typing import List

class DocxSummarizer:
    def __init__(self, api_key: str = "", model: str = "gpt-4o-mini"):
        self.model = model
        self.client = OpenAI(api_key=api_key)

    def docx_to_text_with_tables(self,path):
        doc = Document(path)
        content = []

        para_idx, table_idx = 0, 0  # giữ vị trí paragraph và table

        for element in doc.element.body:
            if element.tag.endswith('p'):  # paragraph
                para = doc.paragraphs[para_idx]
                para_idx += 1
                text = para.text.strip()
                if text:
                    content.append(text)

            elif element.tag.endswith('tbl'):  # table
                table = doc.tables[table_idx]
                table_idx += 1
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    content.append("| " + " | ".join(row_text) + " |")

        return "\n".join(content)


    def summarize_text(self, text: str, max_words: int = 500) -> str:
        """Summarize extracted text using OpenAI."""
        prompt = (
            f"Tóm tắt văn bản sau thành khoảng {max_words} từ. "
            "Yêu cầu: viết rõ ràng, dễ hiểu, bằng tiếng Việt, và trình bày dưới dạng markdown "
            "(sử dụng tiêu đề, gạch đầu dòng nếu phù hợp).\n\n"
            f"Văn bản:\n{text}"
        )

        response = self.client.chat.completions.create(
            model=self.model,
            messages=[
                {"role": "system", "content": "Bạn là một trợ lý AI hữu ích, chuyên tóm tắt văn bản."},
                {"role": "user", "content": prompt},
            ],
            temperature=0.5,
        )

        return response.choices[0].message.content.strip()

    def summarize_docx(self, file_path: str, max_words: int = 200) -> str:
        """End-to-end method: read DOCX and summarize."""
        text = self.docx_to_text_with_tables(file_path)
        return self.summarize_text(text, max_words=max_words)

class DocTagger:
    def __init__(self, api_key: str = "", model: str = "gpt-4o-mini"):
        self.model = model
        self.client = OpenAI(api_key=api_key)

    def docx_to_text_with_tables(self, path):
        doc = Document(path)
        content = []
        para_idx, table_idx = 0, 0

        for element in doc.element.body:
            if element.tag.endswith('p'):
                para = doc.paragraphs[para_idx]
                para_idx += 1
                text = para.text.strip()
                if text:
                    content.append(text)

            elif element.tag.endswith('tbl'):
                table = doc.tables[table_idx]
                table_idx += 1
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    content.append("| " + " | ".join(row_text) + " |")

        return "\n".join(content)

    def classify_tags(self, text: str, list_tags: List[str]) -> List[str]:
        prompt = (
            "Bạn là một hệ thống phân loại văn bản.\n\n"
            f"Danh sách tag cho phép: {', '.join(list_tags)}\n\n"
            "Hãy đọc văn bản dưới đây và chọn tối đa 3 tag phù hợp nhất. "
            "Chỉ trả lời bằng một danh sách tag phân cách bằng dấu phẩy.\n\n"
            f"Văn bản:\n{text}"
        )

        response = self.client.chat.completions.create(
            model=self.model,
            messages=[
                {"role": "system", "content": "Bạn là một trợ lý AI chuyên phân loại văn bản."},
                {"role": "user", "content": prompt},
            ],
            temperature=0,
        )

        tags_text = response.choices[0].message.content.strip()
        tags = [t.strip() for t in tags_text.split(",") if t.strip()]
        valid_tags = [t for t in tags if t in list_tags]
        return valid_tags[:3]

    def classify_docx(self, file_path: str, list_tags: List[str]) -> List[str]:
        text = self.docx_to_text_with_tables(file_path)
        return self.classify_tags(text, list_tags)
